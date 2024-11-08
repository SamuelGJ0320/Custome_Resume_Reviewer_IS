from django.http import JsonResponse, HttpResponse
from django.shortcuts import render
import json
import openai
from dotenv import load_dotenv
import os
from .models import ImprovedCV
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.units import inch


# Carga las API keys y otros valores de entorno
load_dotenv('api_keys_1.env')
openai.api_key = os.getenv('openai_apikey')

def custom_resume_view(request):
    if request.method == 'POST':
        return mejorar_cv(request)  # Llama a la función de mejorar_cv
    return render(request, 'custome_resume.html')

def mejorar_cv(request):
    # Obtener el texto del CV y de la vacante
    cv_text = request.POST.get('cvText')
    vacancy_text = request.POST.get('vacancy')
    output_format = request.POST.get('outputFormat', 'text')

    # Procesar el archivo PDF o DOCX si se cargó uno
    if 'cvFile' in request.FILES:
        uploaded_file = request.FILES['cvFile']
        if uploaded_file.name.endswith('.pdf'):
            cv_text = extract_text_from_pdf(uploaded_file)
        elif uploaded_file.name.endswith('.docx'):
            cv_text = extract_text_from_docx(uploaded_file)

    # Crear el prompt para la IA
    prompt = (
        f"Aquí está la descripción de una vacante: {vacancy_text}\n\n"
        f"Este es el CV actual:\n{cv_text}\n\n"
        f"""
            A continuación te proporcionaré el texto de un currículum vitae y una vacante. 

            Quiero que edites el CV para que resaltes las habilidades y experiencias que son relevantes para la vacante, 
            eliminando información irrelevante que no se ajuste al puesto. 

            No debes inventar información ni eliminar datos personales como el nombre, correo electrónico o número de teléfono. 
            Tampoco incluyas comentarios adicionales ni ningún mensaje como "Aquí está tu CV mejorado"; 
            solo necesito el texto del currículum actualizado.
        """
    )

    # Llamar a la API de OpenAI para generar el nuevo CV
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=1024,
        temperature=0.7,
    )

    # Obtener el texto del CV mejorado
    new_cv = response['choices'][0]['message']['content'].strip()

    # Guardar en la base de datos
    improved_cv_record = ImprovedCV(
        original_cv=cv_text,
        vacancy_description=vacancy_text,
        improved_cv=new_cv
    )
    improved_cv_record.save()

    # Generar la respuesta según el formato de salida seleccionado
    if output_format == 'docx':
        response = generate_docx_response(new_cv)
    elif output_format == 'pdf':
        response = generate_pdf_response(new_cv)
    else:  # Texto
        response = render(request, 'custome_resume.html', {
            'cvText': cv_text,
            'vacancy': vacancy_text,
            'newCv': new_cv
        })
    return response

def extract_text_from_pdf(pdf_file):
    """Extrae texto de un archivo PDF."""
    pdf_reader = PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(docx_file):
    """Extrae texto de un archivo DOCX."""
    doc = Document(docx_file)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    return text

def generate_docx_response(text):
    """Genera una respuesta en formato DOCX."""
    doc = Document()
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=mejorado_cv.docx'
    return response

def generate_pdf_response(text):
    """Genera una respuesta en formato PDF usando ReportLab con ajuste de estilo y ajuste de línea."""
    from textwrap import wrap
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y_position = height - inch  # Posición inicial para el contenido

    # Margen y ancho máximo de línea
    left_margin = 1 * inch
    right_margin = width - inch
    max_width = right_margin - left_margin  # Ancho disponible para el texto

    # Estilos de fuente
    p.setFont("Helvetica", 12)

    # Divide el texto en líneas según el formato Markdown
    lines = text.splitlines()

    for line in lines:
        # Estilos para encabezados
        if line.startswith("## "):  # Encabezado de segundo nivel
            p.setFont("Helvetica-Bold", 12)
            line = line[3:]
            y_position -= 15  # Espacio adicional antes del encabezado
        elif line.startswith("# "):  # Encabezado de primer nivel
            p.setFont("Helvetica-Bold", 14)
            line = line[2:]
            y_position -= 20
        elif line.startswith("- "):  # Elemento de lista
            p.setFont("Helvetica", 12)
            line = f"• {line[2:]}"  # Cambia el guion por un punto de viñeta
        elif "**" in line:  # Texto en negrita
            line = line.replace("**", "")  # Elimina los asteriscos dobles
            p.setFont("Helvetica-Bold", 12)
        else:
            p.setFont("Helvetica", 12)  # Fuente normal para el resto del texto

        # Ajuste para nueva página si el espacio no es suficiente
        if y_position <= inch:
            p.showPage()
            p.setFont("Helvetica", 12)
            y_position = height - inch

        # Ajuste de línea con ancho máximo
        wrapped_lines = wrap(line, width=int(max_width / 6))  # Ajusta el ancho dividiendo por un factor

        # Dibujar cada línea envuelta
        for wrapped_line in wrapped_lines:
            if y_position <= inch:  # Nueva página si no hay espacio
                p.showPage()
                p.setFont("Helvetica", 12)
                y_position = height - inch

            p.drawString(left_margin, y_position, wrapped_line)
            y_position -= 14  # Espacio entre las líneas

    p.save()
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename=mejorado_cv.pdf'
    return response

# pagina de recomendaciones

def fetch_recommendations(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        job_title = data.get("job", "")

        # Llama a OpenAI para obtener recomendaciones
        prompt = f"Proporciona las últimas tendencias de contratación para el puesto de {job_title}."
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=300,
            temperature=0.7,
        )
        recommendations = response['choices'][0]['message']['content'].strip().split('\n')

        return JsonResponse({"recommendations": recommendations})
    return JsonResponse({"error": "Invalid request"}, status=400)
